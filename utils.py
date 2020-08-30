import json
from copy import deepcopy
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement

# reorders a list of elements "x", from newest to latest, in base of a parameter "keyword"
def reorder(x, keyword): return sorted(x, key=lambda k: k.get(keyword, 0), reverse=True)

# adds/sets the attribute "k" to "v", for each element of the list "x"
def setattribute(x, k, v):
	for e in x: e[k] = v
	return x

# starting from a docx document "doc", deletes all unrequired tables (tables with no elements in the cv), defined as all the tables expect used ones (defined by their index through the "usedtableindex" list), and optionally replaces the deleted tables content with text "t"
def deletetables(doc, usedtableindexes, t=''):
	for i in range(0, len(doc.tables)):
		if i in usedtableindexes: continue
		newpar = Paragraph(OxmlElement("w:p"), doc.tables[i])
		newpar.add_run(t)
		newpar_space = Paragraph(OxmlElement("w:p"), doc.tables[i])
		newpar_space.add_run('')
		doc.tables[i]._element.clear()
		doc.tables[i]._element.append(newpar_space._element)
		doc.tables[i]._element.append(newpar._element)

# loads a json list object from a txt file
def loadfromtxtfile(filename):
	f = open(filename, 'r')
	res = []
	for el in f.read().split('\n\n'):
		v = {}
		for row in el.split('\n'):
			if len(row) > 0 and row[0] == '#': continue
			if not ':' in row: continue
			i = row.index(':')
			key = row[:i].strip()
			value = row[i+1:].strip()
			v[key] = value
		res.append(v)
	return res

# loads a json list object from a json file
def loadfromjsonfile(filename):
	f = open(filename, 'r')
	return json.loads(f.read())

# clones the table with index "index" on the docx document "doc", by taking the list "l" of contents (from data.py) and the mappings "mappings" (from mappings.py), optionally, replacing/deleting the content on the original (cloned) table
def addtables(doc, index, l, mappings, replaceoriginaltable):
	t = doc.tables[index]
	newtables = []
	for i in range(0, len(l)): newtables.append(deepcopy(t))
	for i in range(0, len(l)):
		newpar = Paragraph(OxmlElement("w:p"), doc.tables[index])
		newpar.add_run('')
		newpar_space = Paragraph(OxmlElement("w:p"), doc.tables[index])
		newpar_space.add_run('')
		obj = l[i]
		newtable = newtables[i]
		if replaceoriginaltable and i == 0: newtable = doc.tables[index]
		for m in mappings:
			try:
				if m['marker']['type'] == 'index': newtable.cell(int(m['row']),0).text = m['text'].replace(m['marker']['name'], str(i+1))
				if m['marker']['type'] == 'field': newtable.cell(int(m['row']),0).text = m['text'].replace(m['marker']['name'], str(obj[m['marker']['field']]))
				if m['marker']['type'] == 'multiple':
					v = ''
					for mm in m['marker']['values']:
						if mm['marker']['type'] == 'field' and obj.get(mm['marker']['field']) != None: v += mm['text'].replace(mm['marker']['name'], str(obj.get(mm['marker']['field'])))+m['marker']['separator']
					newtable.cell(int(m['row']),0).text = m['text'].replace(m['marker']['name'], v)
			except: newtable.cell(int(m['row']),0).text = m['text'].replace(m['marker']['name'], '-')
			if m.get('style') != None:
				if m['style'] == 'bold': newtable.cell(int(m.get('row')),0).paragraphs[0].runs[0].font.bold = True
		if i == 0: newpar._element.append(newpar_space._element)
		if i > 0 or not replaceoriginaltable:
			newpar._element.append(newtable._tbl)
			doc.tables[index]._tbl.append(newpar._element)
			doc.tables[index]._tbl.append(newtable._tbl)
