from docx import Document

from utils import addtables, deletetables, filterlist
from mappings import *

# should the original table be deleted/replaced with content or kept as reference?
REPLACEORIGINALTABLE = True

# should unused tables be deleted or not?
DELETEUNUSEDTABLES = True
# replacement text when unused tables are deleted
DELETEUNUSEDTABLES_TEXT = 'Nessun elemento in questa categoria.'

# years to exclude
# list of strings of excluded years (empty list "[]"" to ignore)
EXCLUDED_YEARS = []
# field/attribute name to consider (if the specified attribute is not present, the element will be included in your cv)
EXCLUDED_YEARS_FIELD = 'year'

# field used for ordering
ORDER_FIELD = EXCLUDED_YEARS_FIELD

# loading the template document
doc = Document('input.docx')

# paragraphs to remove
for p in paragraphstoremove:
	if not p['enabled']: continue
	l = p['paragraphsindex']
	if not isinstance(p['paragraphsindex'], list): l = [l]
	for e in l: doc.paragraphs[e].text = ''

# updating headers
for h in headers_mappings:
	if not h['enabled']: continue
	doc.paragraphs[h['paragraphindex']].text = h['text']
	try:
		if h['style'] == 'bold': doc.paragraphs[h['paragraphindex']].runs[0].font.bold = True
	except: pass

# updating content
usedtableindexes = []
currentindex = 1
for p in mappings:
	if not p['enabled']: continue
	if p.get('type') != 'mixed':
		l = filterlist(p['list'], EXCLUDED_YEARS_FIELD, EXCLUDED_YEARS)
		if len(l) <= 0: continue
		res = addtables(currentindex, doc, p['tableindex'], l, p['mappings'], replaceoriginaltable=REPLACEORIGINALTABLE)
		currentindex = res.get('currentindex')
		usedtableindexes.append(p['tableindex'])
	else: # should mix more types into a single list
		mixed_list_byyear = {}
		# building the mixed list of objects
		for m in p.get('mappings'):
			if not m['enabled']: continue
			l = filterlist(m['list'], EXCLUDED_YEARS_FIELD, EXCLUDED_YEARS)
			for e in l:
				y = str(e.get(ORDER_FIELD))
				if mixed_list_byyear.get(y) is None: mixed_list_byyear[y] = []
				mixed_list_byyear[y].append({'element':e, 'mapping':m})
		# using the mixed list
		k = sorted(mixed_list_byyear)
		k.reverse()
		if str(k[0]) == 'None':
			k.append(k[0])
			k = k[1:]
		elementstoappend = []
		for y in k:
			a = mixed_list_byyear.get(y)
			for e in a:
				m = e.get('mapping')
				l = [e.get('element')]
				res = addtables(currentindex, doc, m['tableindex'], l, m['mappings'], replaceoriginaltable=(currentindex == 1), positionindex=p['tableindex'], directlyappend=False)
				currentindex = res.get('currentindex')
				elementstoappend += res.get('elementstoappend')
				#usedtableindexes.append(m['tableindex'])
		usedtableindexes.append(p['tableindex'])
		for e in elementstoappend:
			try: doc.tables[p['tableindex']]._tbl.append(e)
			except: pass

# deleting unused tables, if needed
if DELETEUNUSEDTABLES: deletetables(doc, usedtableindexes, DELETEUNUSEDTABLES_TEXT)

# saving the output document
doc.save('output.docx')

# final feedback
print("Please open output.docx and check it accurately.")